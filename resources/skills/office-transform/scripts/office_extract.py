#!/usr/bin/env python3
"""Extract an anchored region of an Office/PDF file into a NEW file.

The source file is opened read-only and never modified. Anchors address the
document's own structural coordinates (worksheet range, body-level paragraph
ordinal, page number) — see SKILL.md for the anchor JSON shapes.

Format-specific third-party readers are imported lazily, so run this with the
dependency matching the source format, e.g.:

    uv run --with openpyxl python office_extract.py \
        --file /abs/report.xlsx \
        --anchor '{"format":"xlsx","sheet":"Sheet1","range":"A1:C10"}' \
        --out /abs/report-extract.csv

Dependencies by source format: xlsx -> openpyxl, docx -> python-docx, pdf -> pypdf.
"""

import argparse
import csv
import json
import re
import sys
from pathlib import Path

A1_CELL_RE = re.compile(r"^([A-Z]{1,3})([1-9][0-9]*)$")


def fail(message: str) -> "sys.NoReturn":
    print(f"error: {message}", file=sys.stderr)
    raise SystemExit(1)


def column_to_index(letters: str) -> int:
    index = 0
    for char in letters:
        index = index * 26 + (ord(char) - ord("A") + 1)
    return index


def parse_a1_cell(ref: str) -> tuple[int, int]:
    match = A1_CELL_RE.match(ref)
    if not match:
        fail(f"invalid A1 cell reference: {ref!r}")
    return column_to_index(match.group(1)), int(match.group(2))


def parse_a1_range(ref: str) -> tuple[int, int, int, int]:
    """Return (min_col, min_row, max_col, max_row) from 'B2' or 'A1:C10'."""
    parts = ref.split(":")
    if len(parts) > 2:
        fail(f"invalid A1 range: {ref!r}")
    start = parse_a1_cell(parts[0])
    end = parse_a1_cell(parts[-1])
    return (
        min(start[0], end[0]),
        min(start[1], end[1]),
        max(start[0], end[0]),
        max(start[1], end[1]),
    )


def slice_char_range(text: str, char_range) -> str:
    if char_range is None:
        return text
    start, end = int(char_range[0]), int(char_range[1])
    if start > end or start < 0:
        fail(f"invalid charRange: {char_range!r}")
    return text[start:end]


def cell_display(value) -> str:
    if value is None:
        return ""
    return str(value)


def write_markdown_table(rows: list[list[str]], out_path: Path) -> None:
    if not rows:
        fail("selection produced no rows")
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    escaped = [[cell.replace("|", "\\|").replace("\n", " ") for cell in row] for row in normalized]
    lines = ["| " + " | ".join(escaped[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in escaped[1:])
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def extract_xlsx(src: Path, anchor: dict, out_path: Path, out_format: str) -> None:
    try:
        from openpyxl import Workbook, load_workbook
    except ImportError:
        fail("openpyxl is required for xlsx sources — rerun via `uv run --with openpyxl python ...`")

    sheet_name = anchor.get("sheet")
    range_ref = anchor.get("range")
    if not sheet_name or not range_ref:
        fail("xlsx anchor requires 'sheet' and 'range'")

    workbook = load_workbook(src, data_only=True, read_only=True)
    if sheet_name not in workbook.sheetnames:
        fail(f"worksheet not found: {sheet_name!r} (has: {workbook.sheetnames})")
    worksheet = workbook[sheet_name]

    min_col, min_row, max_col, max_row = parse_a1_range(range_ref)
    values = [
        [cell.value for cell in row]
        for row in worksheet.iter_rows(min_row=min_row, max_row=max_row, min_col=min_col, max_col=max_col)
    ]
    workbook.close()

    if out_format == "xlsx":
        derived = Workbook()
        derived_sheet = derived.active
        derived_sheet.title = sheet_name[:31]
        for row in values:
            derived_sheet.append(row)
        derived.save(out_path)
    elif out_format == "csv":
        with out_path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.writer(handle)
            writer.writerows([[cell_display(value) for value in row] for row in values])
    elif out_format == "md":
        write_markdown_table([[cell_display(value) for value in row] for row in values], out_path)
    else:
        fail(f"unsupported output format for xlsx source: {out_format!r} (use xlsx, csv, or md)")


def extract_docx(src: Path, anchor: dict, out_path: Path, out_format: str) -> None:
    try:
        import docx
    except ImportError:
        fail("python-docx is required for docx sources — rerun via `uv run --with python-docx python ...`")

    paragraph_index = anchor.get("paragraph")
    if paragraph_index is None or int(paragraph_index) < 0:
        fail("docx anchor requires a non-negative 'paragraph' ordinal")
    paragraph_index = int(paragraph_index)

    document = docx.Document(str(src))
    paragraphs = document.paragraphs
    if paragraph_index >= len(paragraphs):
        fail(f"paragraph {paragraph_index} out of range (document has {len(paragraphs)} body paragraphs)")
    text = slice_char_range(paragraphs[paragraph_index].text, anchor.get("charRange"))

    if out_format in ("txt", "md"):
        out_path.write_text(text + "\n", encoding="utf-8")
    elif out_format == "docx":
        derived = docx.Document()
        derived.add_paragraph(text)
        derived.save(str(out_path))
    else:
        fail(f"unsupported output format for docx source: {out_format!r} (use txt, md, or docx)")


def extract_pdf(src: Path, anchor: dict, out_path: Path, out_format: str) -> None:
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        fail("pypdf is required for pdf sources — rerun via `uv run --with pypdf python ...`")

    page_number = anchor.get("page")
    if page_number is None or int(page_number) < 1:
        fail("pdf anchor requires a one-based 'page' number")
    page_index = int(page_number) - 1

    reader = PdfReader(str(src))
    if page_index >= len(reader.pages):
        fail(f"page {page_number} out of range (document has {len(reader.pages)} pages)")
    page = reader.pages[page_index]

    if out_format == "pdf":
        writer = PdfWriter()
        writer.add_page(page)
        with out_path.open("wb") as handle:
            writer.write(handle)
    elif out_format in ("txt", "md"):
        text = slice_char_range(page.extract_text() or "", anchor.get("charRange"))
        out_path.write_text(text + "\n", encoding="utf-8")
    else:
        fail(f"unsupported output format for pdf source: {out_format!r} (use pdf, txt, or md)")


EXTRACTORS = {"xlsx": extract_xlsx, "docx": extract_docx, "pdf": extract_pdf}


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--file", required=True, help="absolute path of the source document (read-only)")
    parser.add_argument("--anchor", required=True, help="anchor JSON, e.g. '{\"format\":\"xlsx\",...}'")
    parser.add_argument("--out", required=True, help="absolute path of the NEW file to create; must not exist")
    args = parser.parse_args()

    src = Path(args.file)
    out_path = Path(args.out)
    if not src.is_file():
        fail(f"source file not found: {src}")
    if out_path.resolve() == src.resolve():
        fail("output path must differ from the source file — the source is never modified")
    if out_path.exists():
        fail(f"output path already exists: {out_path} — pick a fresh name instead of overwriting")

    try:
        anchor = json.loads(args.anchor)
    except json.JSONDecodeError as error:
        fail(f"anchor is not valid JSON: {error}")
    anchor_format = anchor.get("format")
    extractor = EXTRACTORS.get(anchor_format)
    if extractor is None:
        fail(f"unsupported anchor format: {anchor_format!r} (use xlsx, docx, or pdf)")

    out_format = out_path.suffix.lstrip(".").lower()
    if not out_format:
        fail("output path needs an extension so the output format can be inferred")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    extractor(src, anchor, out_path, out_format)
    print(str(out_path))


if __name__ == "__main__":
    main()
